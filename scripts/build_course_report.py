# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
OUT_DOCX = OUT_DIR / "course_report_gost_7_32_2017.docx"
FIG_DIR = OUT_DIR / "figures"


TITLE = "Обучение модели глубокого обучения для синтеза рукописных траекторий"
SUBTITLE = "Курсовая работа"
SOURCE_ARTICLE = (
    "Graves A. Generating Sequences With Recurrent Neural Networks. "
    "arXiv:1308.0850v5, 2014."
)


ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(100, 100, 100)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.font.size = Pt(9)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(1.25)

    for style_name, size, bold in [
        ("Title", 20, True),
        ("Subtitle", 14, False),
        ("Heading 1", 16, True),
        ("Heading 2", 15, True),
        ("Heading 3", 14, True),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = ACCENT if style_name.startswith("Heading") else RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(10 if style_name.startswith("Heading") else 0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.2

    add_page_number(section.footer.paragraphs[0])
    return doc


def add_plain_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    for run in p.runs:
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = MUTED
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "EAF2F8")
        set_cell_margins(hdr[i], 120, 120, 120, 120)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            p.paragraph_format.first_line_indent = Cm(0)
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            set_cell_margins(cells[i], 100, 120, 100, 120)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.size = Pt(11)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = width
    doc.add_paragraph()
    return table


def create_architecture_figure():
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    path = FIG_DIR / "architecture.png"

    img = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(img)

    def font(size, bold=False):
        candidates = [
            r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
            r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        ]
        for candidate in candidates:
            try:
                return ImageFont.truetype(candidate, size)
            except OSError:
                continue
        return ImageFont.load_default()

    regular = font(34)
    small = font(30)
    fill = "#EAF2F8"
    stroke = "#1F4E79"
    arrow = "#555555"

    def box(x, y, w, h, label):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=18, fill=fill, outline=stroke, width=4)
        lines = label.split("\n")
        line_heights = [draw.textbbox((0, 0), line, font=regular)[3] for line in lines]
        total_h = sum(line_heights) + 10 * (len(lines) - 1)
        cy = y + (h - total_h) / 2
        for line, lh in zip(lines, line_heights):
            bbox = draw.textbbox((0, 0), line, font=regular)
            draw.text((x + (w - (bbox[2] - bbox[0])) / 2, cy), line, fill="#111111", font=regular)
            cy += lh + 10

    def line(start, end):
        draw.line((start, end), fill=arrow, width=5)
        x1, y1 = start
        x2, y2 = end
        if abs(x2 - x1) >= abs(y2 - y1):
            points = [(x2, y2), (x2 - 22, y2 - 12), (x2 - 22, y2 + 12)]
        else:
            points = [(x2, y2), (x2 - 12, y2 - 22), (x2 + 12, y2 - 22)]
        draw.polygon(points, fill=arrow)

    boxes = [
        (70, 120, 250, 120, "Текст\none-hot"),
        (420, 120, 300, 120, "Window\nattention"),
        (820, 120, 220, 120, "LSTM 1"),
        (1120, 120, 220, 120, "LSTM 2"),
        (1420, 120, 220, 120, "LSTM 3"),
        (790, 430, 330, 120, "MDN\nΔx, Δy"),
        (1240, 430, 300, 120, "Pen state\n0/1"),
    ]
    for item in boxes:
        box(*item)

    line((320, 180), (420, 180))
    line((720, 180), (820, 180))
    line((1040, 180), (1120, 180))
    line((1340, 180), (1420, 180))
    line((1530, 240), (1040, 430))
    line((1530, 240), (1390, 430))

    note = "На каждом шаге модель предсказывает следующее смещение пера и вероятность окончания штриха."
    draw.text((90, 660), note, fill="#333333", font=small)
    img.save(path)
    return path


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run("[Название образовательной организации]").bold = True

    for line in ["[Факультет / кафедра]", "[Дисциплина]"]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph(SUBTITLE.upper())
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)

    p = doc.add_paragraph(TITLE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(20)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = [
        ("Выполнил:", "[ФИО студента], группа [номер группы]"),
        ("Руководитель:", "[ФИО преподавателя]"),
        ("Город:", "[город]"),
        ("Год:", "2026"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    for label, value in meta:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
    doc.add_page_break()


def add_contents(doc):
    doc.add_heading("Содержание", level=1)
    items = [
        "Реферат",
        "Введение",
        "1. Теоретические основы синтеза рукописных траекторий",
        "2. Постановка задачи и используемые данные",
        "3. Архитектура разработанной модели",
        "4. Обучение модели",
        "5. Генерация и оценка качества",
        "6. Результаты, ограничения и направления развития",
        "Заключение",
        "Список использованных источников",
        "Приложение А. Структура программной реализации",
    ]
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()


def add_abstract(doc):
    doc.add_heading("Реферат", level=1)
    add_plain_paragraph(
        doc,
        "Курсовая работа посвящена разработке и обучению модели глубокого обучения для синтеза рукописных траекторий "
        "по заданной текстовой строке. Объектом исследования является процесс генерации онлайн-почерка, представленного "
        "последовательностью координат пера и признаков окончания штриха. Предмет исследования — нейросетевая модель "
        "на основе LSTM, оконного механизма внимания и mixture density output."
    )
    add_plain_paragraph(
        doc,
        "Цель работы состоит в создании программной системы, способной обучаться индивидуальному почерку и генерировать "
        "рукописные варианты входного текста. В ходе работы подготовлен пользовательский датасет, реализованы инструменты "
        "разметки, конвертации, обучения, генерации и визуального сравнения чекпоинтов."
    )
    add_plain_paragraph(
        doc,
        "В результате получена модель, способная воспроизводить характерные признаки почерка автора на коротких словах "
        "и фразах. Установлено, что качество генерации зависит не только от функции потерь, но и от эпохи чекпоинта, "
        "параметра bias, объема датасета и устойчивости attention."
    )
    add_plain_paragraph(
        doc,
        "Ключевые слова: синтез рукописного текста, рукописная траектория, LSTM, MDN, attention, глубокое обучение, "
        "онлайн-почерк, генерация последовательностей."
    )
    doc.add_page_break()


def add_terms(doc):
    doc.add_heading("Обозначения и сокращения", level=1)
    add_table(
        doc,
        ["Обозначение", "Расшифровка"],
        [
            ("LSTM", "Long Short-Term Memory, рекуррентная нейронная сеть с ячейками памяти"),
            ("MDN", "Mixture Density Network, выходной слой, задающий параметры смеси распределений"),
            ("RNN", "Recurrent Neural Network, рекуррентная нейронная сеть"),
            ("BCE", "Binary Cross Entropy, функция потерь для бинарной классификации"),
            ("JSON", "текстовый формат хранения траекторий пера"),
            ("NPZ", "архив массивов NumPy, используемый для объединенного датасета"),
            ("bias", "параметр сэмплирования, влияющий на регулярность и разнообразие генерации"),
        ],
        widths=[Cm(4), Cm(11)],
    )
    doc.add_page_break()


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    arch_path = create_architecture_figure()
    doc = setup_document()

    add_title_page(doc)
    add_abstract(doc)
    add_contents(doc)
    add_terms(doc)

    doc.add_heading("Введение", level=1)
    add_plain_paragraph(
        doc,
        "Синтез рукописного текста является задачей генерации последовательности движений пера по заданной строке. "
        "В отличие от простой отрисовки шрифта, такая система должна учитывать форму букв, порядок штрихов, пробелы, "
        "отрывы пера, наклон, вариативность почерка и устойчивость строки. Практическая ценность задачи связана с "
        "созданием персонализированных рукописных образцов, подготовкой данных для исследований и возможной интеграцией "
        "с пишущими устройствами, например с плоттером."
    )
    add_plain_paragraph(
        doc,
        "Цель курсовой работы — разработать и обучить модель глубокого обучения, способную генерировать рукописные "
        "траектории по входному тексту на основе пользовательского датасета. В работе рассматривается не изображение "
        "букв как статическая картинка, а онлайн-представление письма: последовательность координат пера и признаков "
        "окончания штриха."
    )
    add_plain_paragraph(
        doc,
        "Теоретической основой работы стала статья Алекса Грейвса «Generating Sequences With Recurrent Neural Networks», "
        "в которой показано, как рекуррентные сети LSTM могут использоваться для генерации сложных последовательностей, "
        "включая онлайн-почерк, и как модель можно расширить до синтеза рукописного текста с условием на заданную строку."
    )
    add_bullet(doc, "изучить подход LSTM + mixture density output для генерации координат пера;")
    add_bullet(doc, "подготовить датасет из пар JSON-траекторий и текстовой разметки;")
    add_bullet(doc, "реализовать модель синтеза рукописных траекторий, обучить ее и организовать генерацию вариантов;")
    add_bullet(doc, "проанализировать качество, ограничения и перспективы дальнейшего развития системы.")

    doc.add_heading("1. Теоретические основы синтеза рукописных траекторий", level=1)
    doc.add_heading("1.1. Последовательностная природа задачи", level=2)
    add_plain_paragraph(
        doc,
        "Онлайн-почерк естественно представляется как временной ряд. Каждый следующий элемент зависит от предыдущих "
        "движений пера, текущего состояния написания символа и от того, какая часть входного текста уже была реализована "
        "в траектории. Поэтому задача хорошо согласуется с рекуррентными нейронными сетями, которые обрабатывают данные "
        "по шагам и поддерживают внутреннее состояние."
    )
    add_plain_paragraph(
        doc,
        "В статье Graves используется идея пошагового предсказания: сеть получает историю последовательности и выдает "
        "распределение вероятностей для следующей точки. Новая последовательность получается авторегрессионно: сэмпл "
        "из выходного распределения подается обратно как вход на следующем шаге. Такой подход позволяет получать разные "
        "варианты одной и той же строки, что важно для естественного почерка."
    )

    doc.add_heading("1.2. LSTM и долговременные зависимости", level=2)
    add_plain_paragraph(
        doc,
        "Обычные рекуррентные сети плохо удерживают информацию на длинных интервалах, что приводит к нестабильности "
        "при генерации: небольшая ошибка может накапливаться и уводить последовательность от области реальных данных. "
        "LSTM-сети используют ячейки памяти и управляющие ворота, благодаря чему лучше сохраняют информацию о контексте. "
        "Для рукописного текста это важно: модель должна помнить общий стиль, направление строки и уже написанную часть слова."
    )

    doc.add_heading("1.3. Mixture Density Network для координат пера", level=2)
    add_plain_paragraph(
        doc,
        "Координаты пера являются вещественными величинами, поэтому обычная классификация здесь не подходит. В статье "
        "для онлайн-почерка используется mixture density output: нейронная сеть не предсказывает одну фиксированную точку, "
        "а параметры смеси двумерных нормальных распределений. Это позволяет моделировать неоднозначность движения руки: "
        "после одного и того же состояния допустимо несколько похожих продолжений штриха."
    )
    add_plain_paragraph(
        doc,
        "В настоящем проекте каждая точка представляется как x_t = (Δx, Δy, e), где Δx и Δy — смещение пера относительно "
        "предыдущей точки, а e — бинарный признак окончания штриха. Для Δx и Δy используется смесь двумерных гауссиан, "
        "а для признака e — бинарная классификация с функцией потерь BCE."
    )

    doc.add_heading("1.4. Условный синтез по входному тексту", level=2)
    add_plain_paragraph(
        doc,
        "Чтобы генерировать не просто произвольный почерк, а рукописную запись заданной строки, модель должна связывать "
        "траекторию пера с символами текста. В статье Graves для этого используется window layer — мягкий механизм "
        "внимания, который динамически выбирает участок текстовой последовательности. Модель постепенно сдвигает это "
        "окно по символам и тем самым учится решать, какой символ писать в текущий момент."
    )
    add_plain_paragraph(
        doc,
        "В проекте применена близкая идея: символы текста кодируются one-hot векторами, после чего оконный механизм "
        "формирует контекстный вектор для LSTM-декодера. Если окно внимания движется слишком быстро или застревает, "
        "генерация может давать пропуски, лишние штрихи или портить конец фразы."
    )

    doc.add_heading("2. Постановка задачи и используемые данные", level=1)
    add_plain_paragraph(
        doc,
        "Задача формулируется следующим образом: по входной строке на русском языке необходимо сгенерировать "
        "последовательность точек, которая при соединении штрихов визуально соответствует рукописному написанию этой строки. "
        "Система ориентирована на индивидуальный почерк, поэтому обучающий набор собирался вручную для одного автора."
    )
    add_table(
        doc,
        ["Параметр", "Значение в проекте"],
        [
            ("Формат траектории", "JSON: список точек [x, y, pen]"),
            ("Формат разметки", "TXT: строка, соответствующая написанной траектории"),
            ("Объединенный датасет", "dataset/all_trajectories.npz"),
            ("Размер актуального набора", "432 пары JSON/TXT"),
            ("Тип данных", "онлайн-почерк: координаты пера и отрывы"),
            ("Язык", "русский алфавит, пробелы, цифры и знаки препинания"),
        ],
        widths=[Cm(5), Cm(10)],
    )
    add_plain_paragraph(
        doc,
        "Подготовка данных включает несколько этапов: запись траектории, ручное сопоставление с текстовой строкой, "
        "проверку корректности разметки и конвертацию набора в NPZ. Для проверки качества датасета были реализованы "
        "вспомогательные интерфейсы просмотра и редактирования подписей."
    )

    doc.add_heading("3. Архитектура разработанной модели", level=1)
    add_plain_paragraph(
        doc,
        "Разработанная модель повторяет ключевые идеи статьи Graves, но адаптирована к небольшому пользовательскому "
        "датасету русского почерка. На вход поступают символы строки и предыдущий шаг пера. На выходе модель выдает "
        "параметры распределения следующего смещения пера и логит признака окончания штриха."
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(arch_path), width=Inches(5.9))
    add_caption(doc, "Рисунок 1 — Упрощенная схема модели синтеза рукописных траекторий")

    add_table(
        doc,
        ["Компонент", "Назначение"],
        [
            ("TextEmbedding", "one-hot кодирование символов входной строки"),
            ("Window layer", "мягкое внимание по тексту, определяющее текущий символ"),
            ("LSTM-декодер", "учет истории движений пера и контекста написания"),
            ("MDN-выход", "параметры смеси распределений для Δx и Δy"),
            ("Pen output", "вероятность окончания текущего штриха"),
            ("Bias при генерации", "управление балансом между разнообразием и читаемостью"),
        ],
        widths=[Cm(4.5), Cm(10.5)],
    )
    add_plain_paragraph(
        doc,
        "В текущей конфигурации используются три LSTM-слоя по 400 скрытых единиц, 20 компонент смеси для координат "
        "и 10 компонент оконного механизма внимания. Такая конфигурация близка к параметрам, описанным в статье для "
        "задачи синтеза онлайн-почерка."
    )

    doc.add_heading("4. Обучение модели", level=1)
    add_plain_paragraph(
        doc,
        "Обучение выполняется в режиме teacher forcing: на каждом шаге модель получает истинную предыдущую точку из "
        "датасета и предсказывает следующую. Это ускоряет обучение, однако создает отличие от режима генерации, где "
        "предыдущая точка уже является результатом самой модели. Поэтому низкое значение функции потерь не всегда "
        "означает идеальное качество свободной генерации."
    )
    add_table(
        doc,
        ["Параметр", "Используемое значение"],
        [
            ("Оптимизатор", "RMSprop"),
            ("Learning rate", "0.0001"),
            ("Batch size на GTX 1660", "2"),
            ("Grad accumulation", "2"),
            ("Функция потерь координат", "negative log likelihood MDN"),
            ("Функция потерь pen", "binary cross entropy"),
            ("Дополнительная регуляризация attention", "kappa progress loss"),
            ("Сохранение чекпоинтов", "каждые 5 эпох"),
        ],
        widths=[Cm(6), Cm(9)],
    )
    add_plain_paragraph(
        doc,
        "Практически важной частью обучения стала диагностика внимания. На ранних этапах наблюдалась проблема: окно "
        "attention могло слишком быстро проходить по тексту, из-за чего конец строки превращался в короткие штрихи или "
        "точки. Для стабилизации были добавлены начальное смещение kappa, маскирование padding-символов и дополнительный "
        "штраф за некорректный темп продвижения внимания."
    )

    doc.add_heading("5. Генерация и оценка качества", level=1)
    add_plain_paragraph(
        doc,
        "Генерация выполняется авторегрессионно: модель сэмплирует следующее смещение пера из MDN-распределения, затем "
        "использует полученную точку как вход на следующем шаге. Из-за стохастичности один и тот же текст может давать "
        "несколько различных вариантов. Это соответствует природе рукописного текста, но требует визуального отбора "
        "лучших сэмплов."
    )
    add_plain_paragraph(
        doc,
        "Для управления качеством используется параметр bias. При увеличении bias распределение становится более "
        "концентрированным: варианты обычно становятся более регулярными и читаемыми, но теряют часть естественной "
        "вариативности. При слишком большом bias модель может схлопываться в однотипные движения, а при слишком маленьком "
        "— давать больше хаотичных штрихов."
    )
    add_bullet(doc, "читаемость букв и слов;")
    add_bullet(doc, "похожесть на индивидуальный почерк;")
    add_bullet(doc, "правильность числа букв и отсутствие пропусков;")
    add_bullet(doc, "качество пробелов и переходов между словами;")
    add_bullet(doc, "устойчивость горизонтали строки;")
    add_bullet(doc, "корректное завершение фразы без мусорных точек.")

    doc.add_heading("6. Результаты, ограничения и направления развития", level=1)
    add_plain_paragraph(
        doc,
        "На текущем этапе модель уже способна воспроизводить характерные черты пользовательского почерка. На отдельных "
        "коротких фразах получаются варианты, визуально близкие к реальному написанию. При этом качество остается "
        "нестабильным: удачные и неудачные варианты могут появляться для одного и того же текста, поэтому для практического "
        "использования полезно генерировать несколько сэмплов и выбирать лучший."
    )
    add_table(
        doc,
        ["Наблюдение", "Интерпретация"],
        [
            ("После 320 эпох улучшения стали менее заметны", "вероятно, модель приблизилась к локальному оптимуму на текущем датасете"),
            ("Некоторые поздние чекпоинты визуально хуже", "loss при teacher forcing не полностью отражает качество свободной генерации"),
            ("Длинные строки чаще портятся к концу", "ошибки attention и авторегрессионные ошибки накапливаются"),
            ("Горизонталь иногда уходит вверх или вниз", "ошибка в Δy накапливается по шагам"),
            ("Лучшие результаты зависят от bias", "для разных эпох оптимальная степень сэмплирования может отличаться"),
        ],
        widths=[Cm(6), Cm(9)],
    )
    add_plain_paragraph(
        doc,
        "Главное ограничение текущей версии — обучение модели фактически с нуля на почерке одного человека. Для курсовой "
        "работы это приемлемо и позволяет продемонстрировать полный цикл исследования. Для продуктовой системы потребовалась "
        "бы другая стратегия: большая предобученная модель, быстрое дообучение на стиле пользователя или style encoder, "
        "который извлекает вектор почерка из нескольких образцов."
    )

    doc.add_heading("Заключение", level=1)
    add_plain_paragraph(
        doc,
        "В ходе работы была разработана система синтеза рукописных траекторий по входной текстовой строке. Реализация "
        "основана на рекуррентной LSTM-модели с оконным механизмом внимания и mixture density output. Были подготовлены "
        "инструменты для разметки данных, конвертации датасета, обучения, генерации нескольких вариантов и визуального "
        "сравнения чекпоинтов по эпохам и bias."
    )
    add_plain_paragraph(
        doc,
        "Полученные результаты показывают, что выбранная архитектура способна обучаться индивидуальному почерку и "
        "генерировать читаемые рукописные фразы. При этом качество зависит от объема и баланса датасета, стабильности "
        "механизма внимания и параметров сэмплирования. Дальнейшее развитие проекта может быть связано с увеличением "
        "датасета, введением автоматических метрик качества, адаптацией модели к плоттеру и переходом к предобученной "
        "модели с быстрой персонализацией стиля."
    )

    doc.add_heading("Список использованных источников", level=1)
    add_numbered(doc, "ГОСТ 7.32-2017. Система стандартов по информации, библиотечному и издательскому делу. Отчет о научно-исследовательской работе. Структура и правила оформления.")
    add_numbered(doc, "ГОСТ Р 2.105-2019. Единая система конструкторской документации. Общие требования к текстовым документам.")
    add_numbered(doc, SOURCE_ARTICLE)
    add_numbered(doc, "Bishop C. M. Mixture Density Networks. Technical Report, 1994.")
    add_numbered(doc, "Hochreiter S., Schmidhuber J. Long Short-Term Memory. Neural Computation, 1997.")
    add_numbered(doc, "IAM On-Line Handwriting Database: набор данных онлайн-почерка, использованный в исходной статье Graves.")

    doc.add_page_break()
    doc.add_heading("Приложение А. Структура программной реализации", level=1)
    add_table(
        doc,
        ["Файл", "Назначение"],
        [
            ("src/model.py", "описание LSTM-декодера, window layer и MDN-выхода"),
            ("src/train.py", "одна эпоха обучения, функции потерь, collate_fn"),
            ("src/run_training.py", "запуск обучения, checkpoint resume, сохранение моделей"),
            ("src/generate.py", "авторегрессионная генерация траекторий"),
            ("src/run_generate.py", "генерация нескольких вариантов одной строки"),
            ("src/compare_epochs_bias.py", "визуальное сравнение эпох и bias"),
            ("dataset/converter.py", "сборка JSON/TXT в единый NPZ-датасет"),
            ("label_trajectories.py", "интерфейс разметки и проверки траекторий"),
        ],
        widths=[Cm(5), Cm(10)],
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_report()
